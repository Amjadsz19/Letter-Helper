import streamlit as st
from docx import Document, enum
from bs4 import BeautifulSoup
from selenium import webdriver
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

import docx
import shutil
import io
import convert_numbers
import requests

st.set_page_config(page_title="Letter Builder" , layout="wide")

def ArabicToEnglish(paragraph):

    for word in paragraph:
        for letter in word:
            if letter.isdigit():
                paragraph = paragraph.replace(letter, convert_numbers.english_to_arabic(letter))
    


    return paragraph

#-----Define varibles-----
Topic = ""
Title = ""
Receiver = ""
Letter = ""
MainObject = ""

#---------Design the website-------

with open("Designing.css") as source_des:
    st.markdown(f"<style>{source_des.read()}</style>", unsafe_allow_html=True)

#---------Header section----------
with st.container():
    left_column, Center_column, right_column = st.columns(3)
    with Center_column:
        st.title("مساعد إنشاء الخطابات")
        st.text("تم بناء هذا الموقع للمساعدة في إنشاء خطابات سليمة نحوياَ ولغوياَ")
        st.text("وكتابتها في الملف المخصص للخطاب")

    with st.container():
        st.write("---")
        left_column, Center_column, right_column = st.columns(3)
        with Center_column:
            st.header("أبدأ بإنشاء خطابك")
            st.write("##")



with st.container():
    left_column,Center_column, right_column = st.columns(3)
    with Center_column:
        st.markdown(
            """
            <style>

                div[data-testid="column"]:nth-of-type(1)
                {
                    text-align: end;
                } 
            </style>
            """, unsafe_allow_html=True
        )
        st.title("أدخل عنوان الخطاب")
        st.text("*للتوضيح عنوان الخطاب سيكون اسم الملف")
        Topic = st.text_input("")

with st.container():
    st.write("---")
    left_column,Center_column, right_column = st.columns(3)
    with Center_column:

        st.write("##")
        st.title("أدخل موضوع الخطاب")
        Title = st.text_input(" ")

with st.container():
    left_column,Center_column, right_column = st.columns(3)

    with Center_column:
        st.write("##")
        st.title("أدخل مستلم الخطاب")

        Receiver = st.text_input("  ")

with st.container():
    left_column,Center_column, right_column = st.columns(3)

    with Center_column:
        st.write("##")
        st.title("أدخل الخطاب")

        Letter = st.text_area("  ")

    if(st.button("تحقق من صحة الخطاب")):
        if (Letter != ""):
            url = 'https://corrector.app/ar/'

            # open browser
            driver = webdriver.Chrome()

            # load page
            driver.get(url)

            # find field
            driver.implicitly_wait(10)

            frame_0 = driver.find_element(By.ID, 'checktext_ifr')
            driver.switch_to.frame(frame_0)

            item = driver.find_element(By.ID, 'tinymce')
            item1 = item.find_element(By.TAG_NAME, 'p')

            # put text
            item1.send_keys(MainObject)

            driver.switch_to.default_content()

            # find button
            item = driver.find_element(By.ID, 'correct')

            # click button
            driver.execute_script("arguments[0].click();", item)

            frame_0 = driver.find_element(By.ID, 'checktext_ifr')
            driver.switch_to.frame(frame_0)

            # find all errors
            all_answers = driver.find_elements(By.CLASS_NAME, 'hiddenSpellError')
            all_answers1 = driver.find_elements(By.CLASS_NAME, 'hiddenGrammarError')
            st.write("##")

            st.write("أخطاء إملائية:")
            for answer in all_answers:
                st.write(answer.text)
            st.write("##")
            st.write("أخطاء نحوية\n")
            for answer in all_answers1:
                st.write(answer.text)

            driver.quit()

        else:
            st.error("قم بإدخال الخطاب أولاً")




    if (st.button("أنشئ الخطاب")):
        if (Topic == "" or Title == "" or Receiver == "" or Letter == ""):
            st.error("قم بملئ الفراغات قبل الضغط على إنشاء")

        else:
            # Create a file that has the same format of the original to write in
            Topic = ArabicToEnglish(Topic)

            Topic = Topic + ".docx"
            shutil.copyfile('letter.docx', Topic)

            document = Document(Topic)

            # Start writing the paragraphs and format them
            paragraph = document.add_paragraph()
            rtlstyle = document.styles.add_style('rtl', enum.style.WD_STYLE_TYPE.PARAGRAPH)
            mystyle = document.styles.add_style('mystyle', enum.style.WD_STYLE_TYPE.CHARACTER)
            my_font = mystyle.font
            my_font.bold = False
            my_font.name = 'AL-Mohanad'
            my_font.size = docx.shared.Pt(15)
            rtlstyle.font.rtl = True
            
            Title1 = "الموضوع: " + Title
            Title1 = ArabicToEnglish(Title1)
            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style = rtlstyle

            run = para.add_run(Title1)
            # Add some formatting to the run
            run.font.bold = False

            run.font.name = 'AL-Mohanad'
            run.font.size = docx.shared.Pt(15)
            
            
            Receiver = ArabicToEnglish(Receiver)
            Greet = "\n\nالمكرم/ "
            Greet = Greet + Receiver
            Greet = Greet + "                      يحفظه الله"
            Greet = Greet + u"\nالسلام عليكم ورحمة الله وبركاته.."
            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            para.style = rtlstyle
            
            run = para.add_run(Greet)
            
            run.font.bold = True
            
            # # Add some formatting to the run
            run.font.name = 'AL-Mohanad'
            run.font.size = docx.shared.Pt(20)
            
            
            
            Letter = ArabicToEnglish(Letter)
            MainLetter = "\n     "
            MainLetter = MainLetter + Letter
            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            para.style = rtlstyle

            run = para.add_run(MainLetter)
            run.font.bold = False
            
            run.font.name = 'AL-Mohanad'
            run.font.size= docx.shared.Pt(18)


            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.style = rtlstyle

            run = para.add_run("وتقبلوا تحياتي وتقديري\n")
            run.font.bold = True
            # Add some formatting to the run
            run.font.name = 'AL-Mohanad'
            run.font.size = docx.shared.Pt(20)



            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style = rtlstyle

            run = para.add_run("مدير مكتب التعليم بمحافظة الجبيل \n\n")
            run.font.bold = True
            # Add some formatting to the run
            run.font.name = 'AL-Mohanad'
            run.font.size = docx.shared.Pt(20)


            para = document.add_paragraph()
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style = rtlstyle

            run = para.add_run("عازب بن علي الأحمري                                                             \n\n")
            run.font.bold = True
            # Add some formatting to the run
            run.font.name = 'AL-Mohanad'
            run.font.size= docx.shared.Pt(20)

            document.save(Topic)


            bio = io.BytesIO()
            document.save(bio)
            if document:
                st.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=Topic,
                    mime="docx"
                )





